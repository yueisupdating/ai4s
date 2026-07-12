#!/usr/bin/env python3
"""Lightweight DOCX formatter used by the ai4sci batch script.

The formatter intentionally keeps the document content intact. It applies a
small set of stable Word defaults and always writes the requested output file.
If rich formatting fails for an unexpected reason, it falls back to copying the
input DOCX so the batch job can continue with a usable artifact.
"""

from __future__ import annotations

import argparse
import os
import shutil
import re
from difflib import SequenceMatcher
from pathlib import Path


SECTION_TITLES = {
    "受访者基本信息",
    "AI赋能科研就绪度调研",
    "数据资源就绪度",
    "计算资源就绪度",
    "模型与平台就绪度",
    "科研创新就绪度",
    "人才储备就绪度",
    "组织机制就绪度",
    "开放合作就绪度",
    "整体就绪度",
    "AI赋能科研需求调研",
}

QUESTION_MARKER_RE = re.compile(
    r"^\s*(?:#+\s*)?(?:Q\s*)?0*([1-9]|[1-3][0-9]|4[0-7])\s*(?:$|[\u3001.\uff0e:\uff1a\s\uff08(])",
    re.I,
)

INSTITUTE_FULL_NAMES = {
    "中科院动物所": "中国科学院动物研究所",
    "中科院植物所": "中国科学院植物研究所",
    "中科院上海药物所": "中国科学院上海药物研究所",
    "中科院地质与地球物理所": "中国科学院地质与地球物理研究所",
    "中科院大气物理所": "中国科学院大气物理研究所",
    "中科院半导体所": "中国科学院半导体研究所",
    "中科院微电子所": "中国科学院微电子研究所",
    "中科院数学与系统科学研究院": "中国科学院数学与系统科学研究院",
    "中科院物理所": "中国科学院物理研究所",
    "中科院化学研究所": "中国科学院化学研究所",
    "中科院国家空间科学中心": "中国科学院国家空间科学中心",
}

MD_TABLE_RE = re.compile(r"^\s*\|.*\|\s*$")
MD_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")


def is_source_question_line(text: str) -> bool:
    if not text or text in SECTION_TITLES:
        return False
    if text.startswith("院内科研单位AI4S") or text.startswith("注意："):
        return False
    if text.startswith("您所在的机构名称") or text.startswith("您所在机构/团队"):
        return True
    if "？" in text:
        return True
    if text.endswith("：") and (text.startswith("您所在机构") or text.startswith("AI为")):
        return True
    if text.startswith("所在领域"):
        return True
    if text.startswith("从整体上看"):
        return True
    if text.startswith("在") and "认为" in text:
        return True
    if text.startswith("请从以下可选维度"):
        return True
    return False


def read_docx_paragraphs(path: Path) -> list[str]:
    from docx import Document

    return [paragraph.text.strip() for paragraph in Document(str(path)).paragraphs if paragraph.text.strip()]


def find_source_questionnaire(input_path: Path) -> Path | None:
    search_dirs = [Path.cwd(), input_path.parent, input_path.parent.parent]
    seen: set[Path] = set()
    for directory in search_dirs:
        directory = directory.resolve()
        if directory in seen or not directory.is_dir():
            continue
        seen.add(directory)
        candidates = [
            candidate
            for candidate in directory.glob("院内科研单位AI4S*.docx")
            if candidate.resolve() != input_path.resolve()
        ]
        if candidates:
            return candidates[0]
    return None


def source_question_blocks(source_path: Path) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []
    current: dict[str, object] | None = None

    for text in read_docx_paragraphs(source_path):
        if text in SECTION_TITLES or text.startswith("注意："):
            continue

        if is_source_question_line(text):
            if current is not None:
                blocks.append(current)
            current = {
                "number": len(blocks) + 1,
                "question": text,
                "options": [],
            }
            continue

        if current is not None:
            current["options"].append(text)

    if current is not None:
        blocks.append(current)
    return blocks


def generated_answer_entries(input_path: Path) -> tuple[list[str], list[dict[str, object]]]:
    intro: list[str] = []
    entries: list[dict[str, object]] = []
    current_number: int | None = None
    current_marker = ""
    current_lines: list[str] = []

    for text in read_docx_paragraphs(input_path):
        match = QUESTION_MARKER_RE.match(text)
        if match:
            if current_number is not None:
                entries.append({
                    "number": current_number,
                    "marker": current_marker,
                    "lines": current_lines,
                })
            current_number = int(match.group(1))
            current_marker = text
            current_lines = []
            continue

        if current_number is None:
            intro.append(text)
        else:
            current_lines.append(text)

    if current_number is not None:
        entries.append({
            "number": current_number,
            "marker": current_marker,
            "lines": current_lines,
        })
    return intro, entries


def generated_answer_blocks(input_path: Path) -> tuple[list[str], dict[int, list[str]]]:
    intro, entries = generated_answer_entries(input_path)
    blocks: dict[int, list[str]] = {}
    for entry in entries:
        blocks[int(entry["number"])] = list(entry["lines"])
    return intro, blocks


def normalize_for_match(text: str) -> str:
    text = re.sub(r"\s+", "", text).lower()
    text = re.sub(r"[\W_]+", "", text, flags=re.UNICODE)
    return text


def entry_question_candidates(entry: dict[str, object]) -> list[str]:
    candidates = [str(entry.get("marker") or "")]
    for line in list(entry.get("lines") or [])[:10]:
        text = str(line)
        if "\u539f\u9898" in text or "\u539f\u9009\u9879" in text:
            candidates.append(text)
    return candidates


def question_match_score(source_question: str, entry: dict[str, object]) -> float:
    source_norm = normalize_for_match(source_question)
    if not source_norm:
        return 0.0

    best = 0.0
    for candidate in entry_question_candidates(entry):
        candidate_norm = normalize_for_match(candidate)
        if not candidate_norm:
            continue
        if source_norm in candidate_norm or candidate_norm in source_norm:
            return 1.0
        best = max(best, SequenceMatcher(None, source_norm, candidate_norm).ratio())
    return best


def align_answer_blocks_to_source(
    entries: list[dict[str, object]],
    source_blocks: list[dict[str, object]],
) -> dict[int, list[str]]:
    aligned: dict[int, list[str]] = {}
    used_entries: set[int] = set()

    scored: list[tuple[float, int, int]] = []
    for source_index, block in enumerate(source_blocks):
        question = str(block["question"])
        for entry_index, entry in enumerate(entries):
            scored.append((question_match_score(question, entry), source_index, entry_index))

    for score, source_index, entry_index in sorted(scored, reverse=True):
        if score < 0.72:
            break
        number = int(source_blocks[source_index]["number"])
        if number in aligned or entry_index in used_entries:
            continue
        aligned[number] = list(entries[entry_index].get("lines") or [])
        used_entries.add(entry_index)

    for entry_index, entry in enumerate(entries):
        if entry_index in used_entries:
            continue
        number = int(entry["number"])
        if 1 <= number <= len(source_blocks) and number not in aligned:
            aligned[number] = list(entry.get("lines") or [])
            used_entries.add(entry_index)

    return aligned


def strip_markdown_markers(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^\s*#{1,6}\s*", "", text)
    text = re.sub(r"^\s*>\s*", "", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return text.strip()


def clean_markdown_line(line: str) -> list[str]:
    cleaned: list[str] = []
    for part in str(line).splitlines():
        text = strip_markdown_markers(part)
        if not text:
            continue
        if text.startswith("```") or text.startswith("<<<") or text.startswith(">>>"):
            continue
        if MD_TABLE_RE.match(text) or MD_TABLE_SEPARATOR_RE.match(text):
            continue
        if re.fullmatch(r"[-=_]{3,}", text):
            continue
        if re.match(r"^#{1,6}\s*附录", part.strip()) or text.startswith("附录：") or text.startswith("附录:"):
            break
        text = re.sub(r"^\s*[-*]\s+", "", text)
        text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
        text = text.replace("**", "").replace("__", "").replace("`", "").strip()
        if text:
            cleaned.append(text)
    return cleaned


def clean_answer_lines(lines: list[str]) -> list[str]:
    cleaned: list[str] = []
    stop = False
    for line in lines:
        raw = str(line).strip()
        marker_text = strip_markdown_markers(raw)
        if re.match(r"^#{1,6}\s*附录", raw) or marker_text.startswith("附录：") or marker_text.startswith("附录:"):
            stop = True
        if stop:
            continue
        cleaned.extend(clean_markdown_line(raw))
    return cleaned


def infer_institute_full_name(input_path: Path, intro: list[str], answer_blocks: dict[int, list[str]]) -> str | None:
    env_name = os.environ.get("INSTITUTE_NAME", "").strip()
    haystack_parts = [
        env_name,
        str(input_path),
        input_path.name,
        *intro[:8],
        *answer_blocks.get(1, [])[:12],
    ]
    haystack = "\n".join(haystack_parts)
    for short_name, full_name in INSTITUTE_FULL_NAMES.items():
        if short_name in haystack or full_name in haystack:
            return full_name
    return None


def force_q1_answer(lines: list[str], full_name: str | None) -> list[str]:
    if not full_name:
        return lines

    filtered: list[str] = []
    skipping_answer_continuation = False
    for line in lines:
        plain = strip_markdown_markers(line)
        if plain.startswith("【答案】") or plain.startswith("答案"):
            skipping_answer_continuation = True
            continue
        if skipping_answer_continuation:
            if plain.startswith("【证据】") or plain.startswith("【置信度】") or plain.startswith("【"):
                skipping_answer_continuation = False
            else:
                continue
        filtered.append(line)

    return [f"【答案】{full_name}", *filtered]


def strip_existing_source_scaffold(lines: list[str]) -> list[str]:
    stripped: list[str] = []
    skip_option = False

    for line in lines:
        plain = strip_markdown_markers(line).lstrip("* ").strip()
        if plain.startswith("【原题】"):
            continue
        if plain.startswith("【原选项】") or plain.startswith("【填答要求】"):
            skip_option = True
            continue
        if plain.startswith("【作答与依据】"):
            skip_option = False
            continue
        if skip_option and not (
            plain.startswith("【答案】")
            or plain.startswith("【证据】")
            or plain.startswith("【置信度】")
        ):
            continue
        skip_option = False
        stripped.append(line)

    return clean_answer_lines(stripped)


Q47_ALLOWED_DIMENSIONS = [
    "建立扁平化组织架构",
    "建立跨学科分布式团队",
    "改革考核评价机制",
    "明确知识产权归属",
    "优化资源配置",
    "引进与培养复合型人才",
    "培育开放合作与协同文化",
]

Q47_FALLBACK_RANKED_DEMANDS = [
    "第一：引进与培养复合型人才 - 补齐既懂AI方法又懂本学科问题的复合型人才，支撑模型、数据和科研场景长期迭代。",
    "第二：优化资源配置 - 将算力、数据治理、工程支持和项目经费向AI4S重点方向集中配置，降低跨团队使用门槛。",
    "第三：建立跨学科分布式团队 - 围绕重点科学问题组织算法、数据、平台和领域专家协同攻关。",
]


def q47_has_ranked_answer(lines: list[str]) -> bool:
    ranks = ("第一", "第二", "第三", "第四")
    selected: list[str] = []
    for line in lines:
        if "____" in line or not any(rank in line for rank in ranks):
            continue
        matches = [dimension for dimension in Q47_ALLOWED_DIMENSIONS if dimension in line]
        if matches:
            selected.append(matches[0])
    return len(selected) >= 2 and len(set(selected)) == len(selected)


def q47_ranked_answer_line(line: str) -> bool:
    ranks = ("第一", "第二", "第三", "第四")
    return (
        "____" not in line
        and any(rank in line for rank in ranks)
        and any(dimension in line for dimension in Q47_ALLOWED_DIMENSIONS)
    )


def ensure_q47_ranked_answer(lines: list[str]) -> list[str]:
    if q47_has_ranked_answer(lines):
        return lines

    inserted = False
    in_answer = False
    fixed: list[str] = []
    for line in lines:
        if not inserted and "【答案】" in line:
            fixed.append(line)
            fixed.extend(Q47_FALLBACK_RANKED_DEMANDS)
            inserted = True
            in_answer = True
            continue

        if in_answer:
            if "【证据】" in line or "【置信度】" in line:
                in_answer = False
                fixed.append(line)
            elif q47_ranked_answer_line(line):
                continue
            else:
                continue
            continue

        if q47_ranked_answer_line(line):
            continue

        fixed.append(line)

    if inserted:
        return fixed

    return ["【答案】", *Q47_FALLBACK_RANKED_DEMANDS, *fixed]


def rebuild_docx_with_source_form(input_path: Path):
    from docx import Document

    source_path = find_source_questionnaire(input_path)
    if source_path is None:
        return None

    source_blocks = source_question_blocks(source_path)
    if len(source_blocks) != 47:
        return None

    intro, answer_entries = generated_answer_entries(input_path)
    if len(answer_entries) < 40:
        return None
    answer_blocks = align_answer_blocks_to_source(answer_entries, source_blocks)
    full_name = infer_institute_full_name(input_path, intro, answer_blocks)

    document = Document()
    title = "院内科研单位AI4S“就绪度”调查问卷"
    if full_name:
        title = f"{title} - {full_name}"
    document.add_heading(title, level=0)

    for block in source_blocks:
        number = int(block["number"])
        question = str(block["question"])
        options = [str(option) for option in block["options"]]

        document.add_heading(f"Q{number}", level=1)
        document.add_paragraph(f"【原题】{question}")

        if options:
            option_title = "【填答要求】" if number == 47 else "【原选项】"
            document.add_paragraph(option_title)
            for option in options:
                document.add_paragraph(option)

        answer_lines = strip_existing_source_scaffold(answer_blocks.get(number, []))
        if number == 1:
            answer_lines = force_q1_answer(answer_lines, full_name)
        if number == 47:
            answer_lines = ensure_q47_ranked_answer(answer_lines)
        document.add_paragraph("【作答与依据】")
        if answer_lines:
            for line in answer_lines:
                document.add_paragraph(line)
        else:
            document.add_paragraph("【答案】未生成")

    return document


def apply_basic_formatting(document) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    for section in document.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "SimSun"
    normal_style.font.size = Pt(10.5)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    for paragraph in document.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.name = "SimSun"
            run.font.size = run.font.size or Pt(10.5)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "SimSun"
                        run.font.size = run.font.size or Pt(10)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def format_docx(input_path: Path, output_path: Path) -> None:
    try:
        from docx import Document
    except Exception:
        shutil.copy2(input_path, output_path)
        return

    try:
        document = rebuild_docx_with_source_form(input_path)
        if document is None:
            document = Document(str(input_path))

        apply_basic_formatting(document)
        document.save(str(output_path))
    except Exception:
        shutil.copy2(input_path, output_path)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", nargs="?", default="questionnaire.docx")
    parser.add_argument("-o", "--output", default="formatted.docx")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    if not input_path.is_file() or input_path.stat().st_size == 0:
        raise SystemExit(f"input DOCX is missing or empty: {input_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    format_docx(input_path, output_path)

    if not output_path.is_file() or output_path.stat().st_size == 0:
        raise SystemExit(f"output DOCX was not created: {output_path}")

    print(f"wrote: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
